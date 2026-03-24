import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(markdown_file, docx_file):
    doc = Document()
    
    # Set Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    if not os.path.exists(markdown_file):
        print(f"Error: {markdown_file} not found")
        return

    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Headings
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('---'):
            doc.add_page_break()
        # Bullet points
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Bold text (simplistic)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        else:
            doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Successfully created {docx_file}")

if __name__ == "__main__":
    src = r"C:\Users\softs\.gemini\antigravity\brain\8615f4f5-9edb-4a75-b823-8185af313717\project_documentation.md"
    dist = r"C:\Users\softs\.gemini\antigravity\brain\8615f4f5-9edb-4a75-b823-8185af313717\KU_Pharmacy_Project_Documentation.docx"
    create_docx(src, dist)
